from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return h

def add_colored_paragraph(doc, text, bold=False, color=None, size=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    return p

def add_issue_table(doc, issues):
    """
    issues: list of dicts with keys: id, title, severity, category, file, description, action
    severity: CRITICAL | HIGH | MEDIUM | LOW | PLANNED
    """
    SEV_COLORS = {
        'CRITICAL': 'C00000',
        'HIGH':     'E36C09',
        'MEDIUM':   'BF9000',
        'LOW':      '375623',
        'PLANNED':  '17375E',
        'INFO':     '404040',
    }
    SEV_BG = {
        'CRITICAL': 'FFCCCC',
        'HIGH':     'FFE4CC',
        'MEDIUM':   'FFF2CC',
        'LOW':      'EBF1DE',
        'PLANNED':  'DCE6F1',
        'INFO':     'F2F2F2',
    }

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    headers = ['#', 'Severity', 'Category', 'Title', 'File(s)']
    header_widths = [Cm(1), Cm(2.5), Cm(3), Cm(6.5), Cm(4)]
    for i, (cell, txt, w) in enumerate(zip(hdr, headers, header_widths)):
        cell.width = w
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_bg(cell, '1F3864')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for issue in issues:
        row = table.add_row()
        cells = row.cells
        widths = [Cm(1), Cm(2.5), Cm(3), Cm(6.5), Cm(4)]
        sev = issue['severity']
        bg = SEV_BG.get(sev, 'FFFFFF')
        fg = SEV_COLORS.get(sev, '000000')

        vals = [issue['id'], sev, issue['category'], issue['title'], issue.get('file', '—')]
        for cell, val, w in zip(cells, vals, widths):
            cell.width = w
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            set_cell_bg(cell, bg)
            if val == sev:
                run.bold = True
                run.font.color.rgb = RGBColor(*bytes.fromhex(fg))

        # Description row (merged)
        desc_row = table.add_row()
        desc_cells = desc_row.cells
        for c in desc_cells:
            set_cell_bg(c, 'F8F8F8')
        # merge all 5 cells
        merged = desc_cells[0].merge(desc_cells[1]).merge(desc_cells[2]).merge(desc_cells[3]).merge(desc_cells[4])
        merged.paragraphs[0].clear()
        p = merged.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.1)
        r1 = p.add_run('Description: ')
        r1.bold = True
        r1.font.size = Pt(8.5)
        r2 = p.add_run(issue['description'])
        r2.font.size = Pt(8.5)
        # Action row
        action_row = table.add_row()
        action_cells = action_row.cells
        for c in action_cells:
            set_cell_bg(c, 'EAEAEA')
        merged2 = action_cells[0].merge(action_cells[1]).merge(action_cells[2]).merge(action_cells[3]).merge(action_cells[4])
        merged2.paragraphs[0].clear()
        p2 = merged2.paragraphs[0]
        p2.paragraph_format.left_indent = Inches(0.1)
        r3 = p2.add_run('Recommended Action: ')
        r3.bold = True
        r3.font.size = Pt(8.5)
        r4 = p2.add_run(issue['action'])
        r4.font.size = Pt(8.5)
        r4.font.italic = True
        # spacer row
        sp = table.add_row()
        for c in sp.cells:
            set_cell_bg(c, 'FFFFFF')
            c.paragraphs[0].clear()
            c.paragraphs[0].add_run(' ').font.size = Pt(4)

    doc.add_paragraph()

# ─── Build Document ──────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Title
title = doc.add_heading('The Assembly — Full Codebase Audit Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

subtitle = doc.add_paragraph('Prepared: March 2026  |  Version: 1.0  |  Classification: Internal Development')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.font.size = Pt(9)

doc.add_paragraph()

# ─── SECTION 1: Executive Summary ────────────────────────────────────────────
add_heading(doc, '1. Executive Summary', 1, '1F3864')
doc.add_paragraph(
    'This report presents a comprehensive audit of The Assembly codebase as of March 2026. '
    'The project is a sophisticated real-time social deduction game built on a Node.js / Express / Socket.IO backend '
    'with a React 19 frontend. It integrates Supabase for persistence, Redis for caching, Stripe for payments, '
    'and is designed for deployment on Google Cloud Run.\n\n'
    'The game engine is functionally rich and architecturally coherent for its current stage. '
    'However, three monolithic files (server.ts, gameEngine.ts, apiRoutes.ts) represent the primary source of technical debt. '
    'The planned refactoring of these files — in conjunction with the pending Resend email integration — '
    'will substantially improve maintainability, testability, and long-term scalability.\n\n'
    'The audit has identified 0 critical security vulnerabilities in the current codebase (previously patched), '
    'a small number of high-priority logical and stability issues, and a larger body of medium and low priority improvements.'
)

# ─── SECTION 2: What's Good ──────────────────────────────────────────────────
add_heading(doc, '2. Strengths', 1, '375623')

add_heading(doc, '2.1 Game Engine Architecture', 2)
doc.add_paragraph(
    'The GameEngine class employs a single canonical enterPhase() method as the sole path for all game state transitions. '
    'This design correctly eliminates race conditions that would otherwise arise from multiple callers mutating phase state independently. '
    'The use of a strict phase enum ("Nominate_Chancellor", "Voting", "Legislative_President", etc.) makes the state machine '
    'auditable and predictable.'
)

add_heading(doc, '2.2 Bayesian Suspicion Model (suspicion.ts)', 2)
doc.add_paragraph(
    'The AI reasoning system is a genuine standout feature. Rather than simple heuristics, it uses a log-odds Bayesian model '
    'to compound evidence from voting outcomes, policy declarations, inconsistency detection, investigation results, and nomination patterns. '
    'This produces believable, non-trivial AI behaviour that improves with game duration — rare in indie-scale social deduction games.'
)

add_heading(doc, '2.3 Security Posture', 2)
doc.add_paragraph(
    'The server demonstrates strong security hygiene:'
)
for item in [
    'JWT with per-user tokenVersion for cryptographic session invalidation on password reset.',
    'Helmet with a strict Content Security Policy including dynamically computed SHA-256 hashes for production scripts.',
    'Token-bucket rate limiting on both socket game actions and chat messages.',
    'SSRF-mitigated proxy route with an explicit hostname allowlist and protocol enforcement.',
    'Cryptographically secure (randomBytes) invite code generation.',
    'Zod runtime schema validation in supabaseService.ts guarding all database reads.',
    'Admin middleware in apiRoutes.ts with re-validation on every privileged endpoint.',
]:
    add_bullet(doc, item)

add_heading(doc, '2.4 Persistence & Resilience', 2)
doc.add_paragraph(
    'The write-through Redis cache (stateClient separate from pub/sub clients) means game state survives server restarts. '
    'The restoreFromRedis() method correctly rehydrates in-progress rooms on startup. '
    'TTL-based expiry (24 hours) prevents Redis from filling up with abandoned room state.'
)

add_heading(doc, '2.5 Post-Match Economy & Progression', 2)
doc.add_paragraph(
    'The XP/ELO/IP/CP reward pipeline is well-structured. ELO computation uses opponent-average rather than fixed constants, '
    'correctly accounting for team-vs-team asymmetry. The configurable global XP/IP multipliers via the admin panel provide '
    'live economy levers without a deployment. The O(log n) XP level cache from the recent optimisation is a clean improvement.'
)

add_heading(doc, '2.6 Personal Agendas & Achievements', 2)
doc.add_paragraph(
    'The personalAgendas.ts module is a well-isolated, pure-function system (evaluate takes only GameState and playerId). '
    'The 21 agenda definitions and 25 achievements provide meaningful meta-progression alongside faction victory. '
    'Agenda noise in AI voting behaviour (e.g., chaos_agent, the_hawk) is a nice detail.'
)

add_heading(doc, '2.7 Structured Logging', 2)
doc.add_paragraph(
    'Pino logging with Cloud Logging-compatible severity labels is production-ready. '
    'The logger is imported universally and used consistently throughout server.ts and the service layer. '
    'Sensitive data (passwords) is not logged.'
)

add_heading(doc, '2.8 Graceful Shutdown', 2)
doc.add_paragraph(
    'SIGTERM / SIGINT handlers emit a serverRestarting event to connected clients before closing, '
    'giving the frontend a chance to display a reconnection prompt. This is a good production practice for Cloud Run rolling deployments.'
)

# ─── SECTION 3: What's Bad ───────────────────────────────────────────────────
add_heading(doc, '3. Weaknesses', 1, 'E36C09')

add_heading(doc, '3.1 Monolithic File Size', 2)
doc.add_paragraph(
    'The three core files exceed maintainable complexity thresholds:'
)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
for i, row_data in enumerate([
    ('File', 'Lines', 'Responsibility'),
    ('server.ts', '1,311', 'HTTP setup, Socket.IO event hub, Stripe webhook, CSP, proxy, graceful shutdown'),
    ('gameEngine.ts', '2,853', 'All game phases, AI logic, ELO, stats, Redis read/write, disconnection'),
    ('apiRoutes.ts', '1,281', 'Auth, profile, shop, Stripe checkout, admin endpoints, email recovery'),
]):
    cells = table.rows[i].cells
    for j, (cell, val) in enumerate(zip(cells, row_data)):
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if i == 0:
            r.bold = True
            set_cell_bg(cell, '1F3864')
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
doc.add_paragraph()

add_heading(doc, '3.2 Mixed Concerns in GameEngine', 2)
doc.add_paragraph(
    'The GameEngine class combines: game phase logic, AI turn scheduling, Redis persistence, ELO calculation, '
    'achievement checking, match result saving, and disconnection handling. This violates the Single Responsibility Principle '
    'and makes unit testing individual subsystems extremely difficult. Each concern should be its own module.'
)

add_heading(doc, '3.3 Dual Rate-Limiting (Redundant)', 2)
doc.add_paragraph(
    'Chat messages are rate-limited twice: once in the socket.use() middleware (token-bucket) and once in the sendMessage '
    'handler (lastChatTimestamp, 1s cooldown). This is not a bug but adds dead code weight. The middleware bucket should be the single authority.'
)

add_heading(doc, '3.4 In-Memory Fallbacks in Production Path', 2)
doc.add_paragraph(
    'When REDIS_URL is not set, the server silently falls back to in-memory room storage with no persistence. '
    'While intentional for development, there is no runtime guard that prevents this configuration from being deployed to production. '
    'A missing REDIS_URL in a Cloud Run secret should produce a hard startup failure, not a silent degradation.'
)

add_heading(doc, '3.5 Stripe Initialised Unconditionally', 2)
doc.add_paragraph(
    'new Stripe(process.env.STRIPE_SECRET_KEY || \'\') is called at module load time, meaning a missing key produces a '
    'Stripe instance initialised with an empty string. The existing warn log is insufficient — the Stripe endpoints should return '
    'a 503 proactively rather than failing at the point of first real API call with a cryptic auth error.'
)

add_heading(doc, '3.6 playAgain Handler in server.ts', 2)
doc.add_paragraph(
    'The playAgain socket handler in server.ts contains ~70 lines of game state reset logic that directly mutates GameState. '
    'This should be a resetRoom() method on GameEngine, consistent with the pattern used elsewhere (startGame, handleLeave, etc.).'
)

add_heading(doc, '3.7 No Integration or End-to-End Tests', 2)
doc.add_paragraph(
    'Vitest is present as a dependency but no test files are visible in the codebase. '
    'The game engine\'s phase-transition logic, ELO computation, suspicion model, and agenda evaluation are all prime candidates '
    'for pure unit tests. Without these, every refactor carries regression risk.'
)

add_heading(doc, '3.8 Audio: OGG Not Supported on Safari / iOS', 2)
doc.add_paragraph(
    'audio.ts documents this limitation but it is unresolved. All Kenney sound effects are served as .ogg files from gamesounds.xyz. '
    'Safari and all iOS browsers do not support OGG natively, resulting in silent failures for a significant user demographic, '
    'especially relevant given the Discord Activity deployment target.'
)

add_heading(doc, '3.9 Missing Pagination on Leaderboard / Match History', 2)
doc.add_paragraph(
    'The leaderboard query in apiRoutes.ts fetches the top N rows without cursor-based pagination. '
    'As the user base grows this will become expensive. Match history similarly loads all records for a user without limits.'
)

add_heading(doc, '3.10 Nodemailer / Gmail SMTP', 2)
doc.add_paragraph(
    'The current password recovery implementation uses Gmail SMTP (Nodemailer). '
    'This is fragile in production: Google can revoke App Password access, impose sending limits, and flag high-volume sends as spam. '
    'This is a known planned migration to Resend and should be prioritised once a custom domain is verified.'
)

# ─── SECTION 4: What's Ugly ──────────────────────────────────────────────────
add_heading(doc, '4. Critical Issues ("The Ugly")', 1, 'C00000')

add_heading(doc, '4.1 settimeout Closures Over Stale State', 2)
doc.add_paragraph(
    'Multiple setTimeout callbacks in gameEngine.ts capture the outer GameState object (s or state) by reference at call time, '
    'then re-fetch via this.rooms.get(roomId) inside the callback. While the re-fetch pattern is mostly correct, '
    'there are at least two locations (AI banter at round start ~L871, AI banter post-declaration ~L1419) that partially '
    'reference the outer closure variable (s / state) after modifying it through the re-fetched st reference. '
    'This can cause chat messages to be appended to a stale snapshot and broadcast incorrectly.'
)

add_heading(doc, '4.2 Indentation / Brace Mismatch in updateUserStats (~L2218)', 2)
doc.add_paragraph(
    'The closing brace for the socket emissions loop at line 2227 indents one level too few — it closes the outer for-loop '
    'instead of the if(player) block. This is currently harmless because the next statement is the end of the function, '
    'but it represents a structural inconsistency that will cause confusion during future edits and could mask a real bug '
    'if the function is extended.'
)

add_heading(doc, '4.3 kickPlayer Uses Wrong ID Field', 2)
doc.add_paragraph(
    'The kickPlayer socket handler (server.ts ~L1099) looks up the target player by p.id === targetSocketId. '
    'Since the stable player identity refactor (playerSessionId), p.id is a persistent UUID — not a socket ID. '
    'The client likely emits the player\'s session UUID, but the variable name targetSocketId is misleading and can cause '
    'lookup failures if the client ever sends the actual socket ID. This should be renamed and the comment clarified.'
)

add_heading(doc, '4.4 Inconclusive Ranked Game Does Not Save Stats', 2)
doc.add_paragraph(
    'When a player leaves a ranked game in progress (leaveRoom or pause timeout), the game phase is set to "GameOver" with '
    'winner=undefined and no call to updateUserStats(). This means the game is erased from history — gamesPlayed is not incremented, '
    'ELO is unchanged, and no match record is written. Players can exploit this by disconnecting to avoid ELO loss.'
)

add_heading(doc, '4.5 Redis Adapter Uses Same Client as State Reads (Possible Blocking)', 2)
doc.add_paragraph(
    'Although redis.ts creates three separate clients (pub, sub, state), the stateClient is a duplicate() of pubClient. '
    'Under extreme load, commands queued for state reads/writes could block behind pub/sub traffic if the underlying TCP '
    'connection backs up. Using truly independent connections (separate createClient() calls) eliminates this risk.'
)

# ─── SECTION 5: Change List ──────────────────────────────────────────────────
add_heading(doc, '5. Prioritised Change List', 1, '1F3864')
doc.add_paragraph(
    'The following table lists all recommended changes ordered by severity. '
    '"PLANNED" items are already on the roadmap and are listed for completeness. '
    '"CRITICAL" items represent bugs or exploit vectors that should be addressed before any public deployment.'
)

issues = [
    # ── CRITICAL ──
    {
        'id': 'C-01',
        'title': 'Inconclusive Games Must Save Stats (Anti-Exploit)',
        'severity': 'CRITICAL',
        'category': 'Game Logic / Economy',
        'file': 'gameEngine.ts (handleLeave, handlePauseTimeout)',
        'description': 'When a ranked game ends as inconclusive (player disconnect / intentional leave), updateUserStats() '
                       'is never called. Players can avoid ELO loss by deliberately disconnecting. The game simply vanishes from '
                       'all records with zero consequence.',
        'action': 'Create a endGameInconclusive() path that calls updateUserStats() with a penalty payload (e.g., treat as a loss '
                  'for the leaver, neutral for those remaining). Write a match record with winReason="inconclusive".',
    },
    {
        'id': 'C-02',
        'title': 'Stale Closure in AI Banter setTimeout Callbacks',
        'severity': 'CRITICAL',
        'category': 'Game Engine Stability',
        'file': 'gameEngine.ts (~L871, ~L1419)',
        'description': 'AI banter setTimeout callbacks partially reference the outer closure variable (s/state) after performing '
                       'a fresh re-fetch (st = this.rooms.get(roomId)). postAIChat() and broadcastState() are then called '
                       'on a mix of the old and new references, risking stale data broadcasting.',
        'action': 'Ensure all code inside setTimeout exclusively uses the re-fetched reference (st). Remove all direct uses '
                  'of the closure variable (s) inside the callback body.',
    },
    # ── HIGH ──
    {
        'id': 'H-01',
        'title': 'Missing Hard Guard for Production Redis Configuration',
        'severity': 'HIGH',
        'category': 'Infrastructure / Reliability',
        'file': 'server.ts, redis.ts',
        'description': 'If REDIS_URL is not configured in production, the server silently uses in-memory storage with no persistence. '
                       'Any crash or rolling deploy wipes all active game rooms. There is no runtime check preventing this misconfiguration.',
        'action': 'Add a startup guard: if (process.env.NODE_ENV === "production" && !isRedisConfigured) { logger.fatal("REDIS_URL is required in production"); process.exit(1); }',
    },
    {
        'id': 'H-02',
        'title': 'playAgain State Reset Should Be a GameEngine Method',
        'severity': 'HIGH',
        'category': 'Architecture / Maintainability',
        'file': 'server.ts (~L966-L1044)',
        'description': 'The playAgain socket handler directly resets ~30 GameState fields inline in server.ts. '
                       'This is a business logic violation — state mutation should be encapsulated in the GameEngine class, '
                       'consistent with all other game actions.',
        'action': 'Move the reset logic to a GameEngine.resetRoom(roomId) method and call it from the socket handler. '
                  'This also enables testing the reset logic in isolation.',
    },
    {
        'id': 'H-03',
        'title': 'kickPlayer Uses Misleading ID Variable Name (Potential Bug)',
        'severity': 'HIGH',
        'category': 'Game Logic / Bug Risk',
        'file': 'server.ts (~L1091)',
        'description': 'The kickPlayer handler receives targetSocketId from the client but now looks up p.id '
                       '(the stable session UUID). The naming inconsistency risks a future regression if the client is ever '
                       'updated to send the socket ID instead.',
        'action': 'Rename the parameter to targetPlayerId. Verify the client emits the player session UUID. '
                  'Add a comment clarifying the identity model.',
    },
    {
        'id': 'H-04',
        'title': 'Stripe 503 Response Must Be Proactive, Not Late-Fail',
        'severity': 'HIGH',
        'category': 'Infrastructure / Reliability',
        'file': 'server.ts, apiRoutes.ts',
        'description': 'Stripe is instantiated unconditionally with an empty string key if STRIPE_SECRET_KEY is absent. '
                       'The failure is deferred to the first real API call, which produces a cryptic Stripe auth error '
                       'rather than a clean 503.',
        'action': 'In registerRoutes(), wrap all Stripe endpoints in a guard that immediately returns res.status(503).json('
                  '{ error: "Payment service not configured." }) if !process.env.STRIPE_SECRET_KEY.',
    },
    {
        'id': 'H-05',
        'title': 'OGG Audio: Silent Failure on Safari / iOS',
        'severity': 'HIGH',
        'category': 'Cross-Platform Compatibility',
        'file': 'src/lib/audio.ts',
        'description': 'All sound effects are served from gamesounds.xyz as .ogg files. Safari and all iOS browsers '
                       '(including Chrome/Firefox on iOS due to WebKit mandate) do not support OGG, resulting in completely '
                       'silent gameplay for this user group.',
        'action': 'Host MP3 equivalents on GCS alongside the existing custom tracks. Update SOUND_PACKS to use MP3 URLs, '
                  'or implement HTMLAudioElement canPlayType() detection and serve the appropriate format.',
    },
    {
        'id': 'H-06',
        'title': 'No Test Coverage Whatsoever',
        'severity': 'HIGH',
        'category': 'Quality Assurance',
        'file': 'All server/** modules',
        'description': 'Vitest is installed but there are zero test files. The game engine contains complex, '
                       'interdependent logic (phase transitions, ELO, suspicion, agenda evaluation) that is currently '
                       'impossible to regression-test during the planned refactor.',
        'action': 'Before beginning the monolith refactor: write unit tests for (1) computeEloChange, (2) all agenda.evaluate() '
                  'functions, (3) getExecutiveAction, (4) enterPhase transition guards. These are pure functions with no I/O dependencies.',
    },
    # ── MEDIUM ──
    {
        'id': 'M-01',
        'title': 'Remove Redundant Chat Rate Limiting',
        'severity': 'MEDIUM',
        'category': 'Code Quality',
        'file': 'server.ts (~L948-L953)',
        'description': 'Chat is rate-limited twice: by the token-bucket socket middleware and by a per-message timestamp check. '
                       'The timestamp check is redundant dead code.',
        'action': 'Remove the lastChatTimestamp guard from the sendMessage handler. Tighten the middleware bucket if needed.',
    },
    {
        'id': 'M-02',
        'title': 'Redis stateClient Should Be Fully Independent',
        'severity': 'MEDIUM',
        'category': 'Infrastructure / Reliability',
        'file': 'server/redis.ts',
        'description': 'stateClient is created via pubClient.duplicate(). Under sustained load, a blocked pub/sub pipeline '
                       'could delay state read/write operations sharing the same underlying connection pool.',
        'action': 'Replace pubClient.duplicate() with a separate createClient() call for stateClient to ensure fully independent connections.',
    },
    {
        'id': 'M-03',
        'title': 'Brace Indentation Bug in updateUserStats Socket Emission Block',
        'severity': 'MEDIUM',
        'category': 'Code Quality / Bug Risk',
        'file': 'gameEngine.ts (~L2218-L2229)',
        'description': 'The closing brace of the if(player) block inside the post-match socket emission loop appears to '
                       'be indented one level too few, suggesting a structural inconsistency that masks the true scope boundary.',
        'action': 'Review and fix indentation of the closing brace. Add an explicit guard ensuring postMatchResult is only '
                  'emitted when player is found.',
    },
    {
        'id': 'M-04',
        'title': 'Add Pagination to Leaderboard and Match History Endpoints',
        'severity': 'MEDIUM',
        'category': 'Performance / Scalability',
        'file': 'server/apiRoutes.ts',
        'description': 'The leaderboard fetches the top N rows and match history fetches all records for a user. '
                       'Both will become expensive as the user base grows.',
        'action': 'Implement cursor-based pagination with a limit/offset or created_at cursor. Expose limit and cursor '
                  'query parameters. Cap at 50 rows per request.',
    },
    {
        'id': 'M-05',
        'title': 'Refactor apiRoutes.ts into Domain-Scoped Routers',
        'severity': 'MEDIUM',
        'category': 'Architecture (Planned Refactor)',
        'file': 'server/apiRoutes.ts (1,281 lines)',
        'description': 'apiRoutes.ts mixes authentication, user profile management, shop/Stripe, admin tooling, '
                       'and email recovery into a single monolithic registerRoutes() function. '
                       'This is the planned refactor already on the roadmap.',
        'action': 'Split into: routes/auth.ts (login, register, password reset), routes/profile.ts (stats, cosmetics), '
                  'routes/shop.ts (Stripe checkout, CP), routes/admin.ts (user management, config, Redis), '
                  'routes/social.ts (friends, leaderboard). Mount via express.Router().',
    },
    {
        'id': 'M-06',
        'title': 'Refactor server.ts into Focused Modules',
        'severity': 'MEDIUM',
        'category': 'Architecture (Planned Refactor)',
        'file': 'server.ts (1,311 lines)',
        'description': 'server.ts combines HTTP/Socket setup, all socket event handlers, the Stripe webhook, '
                       'the SSRF proxy, Vite dev middleware, and graceful shutdown into one file.',
        'action': 'Extract: server/socket/handlers.ts (all socket.on() blocks), server/middleware/security.ts '
                  '(helmet/CORS/CSP), server/routes/webhooks.ts (Stripe webhook), server/routes/proxy.ts. '
                  'Keep server.ts as a thin orchestration entry point only.',
    },
    {
        'id': 'M-07',
        'title': 'Refactor gameEngine.ts into Service Modules',
        'severity': 'MEDIUM',
        'category': 'Architecture (Planned Refactor)',
        'file': 'server/gameEngine.ts (2,853 lines)',
        'description': 'The GameEngine class is responsible for too many distinct concerns. '
                       'This is the most complex planned refactor.',
        'action': 'Extract: services/phaseManager.ts (enterPhase, phase transitions), services/aiService.ts '
                  '(all AI decision logic), services/statsService.ts (ELO, XP, match saving), '
                  'services/redisService.ts (persistence), services/timerService.ts (action timers, pause). '
                  'GameEngine becomes a facade coordinating these services.',
    },
    {
        'id': 'M-08',
        'title': 'Migrate Password Recovery to Resend',
        'severity': 'MEDIUM',
        'category': 'Infrastructure (Planned)',
        'file': 'server/apiRoutes.ts (forgot-password, reset-password)',
        'description': 'The current Nodemailer / Gmail SMTP implementation is fragile and subject to rate limits and '
                       'deliverability issues. Resend is already in the dependencies awaiting DNS verification.',
        'action': 'Once the custom domain DNS TXT record is verified via Resend dashboard: replace Nodemailer with '
                  'the Resend SDK. Use the resend.emails.send() API with a verified from address. '
                  'Add RESEND_API_KEY to the Cloud Run secret manager.',
    },
    # ── LOW ──
    {
        'id': 'L-01',
        'title': 'Warn on Missing Optional Environment Variables at Startup',
        'severity': 'LOW',
        'category': 'Operations / Observability',
        'file': 'server.ts',
        'description': 'Missing STRIPE_WEBHOOK_SECRET, EMAIL_USER, and RESEND_API_KEY are not logged at startup, '
                       'making misconfiguration hard to diagnose until a relevant endpoint is hit.',
        'action': 'Add a startup check block that logs a warning for each missing optional env var at INFO level, '
                  'so the Cloud Run startup logs clearly enumerate what is and isn\'t configured.',
    },
    {
        'id': 'L-02',
        'title': 'sendMessage Handler: Type Is Always "text" (Missing message.type)',
        'severity': 'LOW',
        'category': 'Code Quality',
        'file': 'server.ts (~L961)',
        'description': 'The sendMessage handler pushes messages with no type field. The Message interface includes an '
                       'optional type field ("text" | "round_separator"). All human chat messages should set type: "text" '
                       'explicitly for consistency with the Message interface.',
        'action': 'Add type: "text" to the message object pushed in the sendMessage handler.',
    },
    {
        'id': 'L-03',
        'title': 'Chat Log Truncation: Off-by-One Across Handlers',
        'severity': 'LOW',
        'category': 'Code Quality',
        'file': 'server.ts, gameEngine.ts (multiple locations)',
        'description': 'state.messages is truncated with > 50 (remove when length exceeds 50) in most places but some '
                       'callsites only truncate state.log, leaving messages to grow unboundedly in rare code paths.',
        'action': 'Create a single addMessage(state, msg) helper (analogous to addLog) that always enforces the 50-item cap.',
    },
    {
        'id': 'L-04',
        'title': 'userJoined Friend Notification: Duplicate getFriends Calls',
        'severity': 'LOW',
        'category': 'Performance',
        'file': 'server.ts (userConnected, joinRoom handlers)',
        'description': 'Both userConnected and joinRoom independently call getFriends() and iterate the userSockets map '
                       'to notify online friends. This means a user joining a room fires getFriends() twice.',
        'action': 'Extract a notifyFriendsOnline(userId, roomId?) helper. Call it once from joinRoom and skip the '
                  'duplicate notification in userConnected when a room join immediately follows.',
    },
    {
        'id': 'L-05',
        'title': 'AI Personality "Chaotic" Has No Description in Code',
        'severity': 'LOW',
        'category': 'Maintainability',
        'file': 'server/gameEngine.ts (choosePolicyToDiscard, aiCastVotes)',
        'description': 'The "Chaotic" personality is used in voting and AI decision logic but has no documentation '
                       'comment describing its intent. Its behaviour (random with a 30% base lie rate) is unclear at a glance.',
        'action': 'Add a JSDoc comment to each branch handling "Chaotic" describing its intent. '
                  'Consider adding a personality definition registry similar to AI_BOTS.',
    },
    {
        'id': 'L-06',
        'title': 'updateMediaState Iterates All Rooms on Every Signal',
        'severity': 'LOW',
        'category': 'Performance',
        'file': 'server.ts (~L1154)',
        'description': 'The updateMediaState handler iterates every room in engine.rooms to find the player, '
                       'rather than using socket.rooms to determine which room the socket is already in.',
        'action': 'Use the getRoom() helper pattern already used in other handlers instead of iterating all rooms.',
    },
    # ── INFO / NICE TO HAVE ──
    {
        'id': 'I-01',
        'title': 'Add Environment Variable Schema Validation (e.g., zod or envalid)',
        'severity': 'INFO',
        'category': 'Developer Experience',
        'file': 'server.ts',
        'description': 'Environment variables are scattered across the codebase and accessed with process.env.X || fallback. '
                       'A centralised, validated env schema (using zod or the envalid library) would catch misconfiguration '
                       'at startup and provide type-safe access to env vars.',
        'action': 'Create server/env.ts that exports a validated environment object using zod.parse(). '
                  'Replace all process.env.X references with env.X.',
    },
    {
        'id': 'I-02',
        'title': 'Consider a @the-assembly/types Shared Package',
        'severity': 'INFO',
        'category': 'Architecture / DX',
        'file': 'src/types.ts, src/sharedConstants.ts',
        'description': 'The project already uses a hybrid import pattern (types from src/ inside server/ files). '
                       'As the monolith is split, a shared internal package would formalise this boundary.',
        'action': 'Create a packages/types directory with its own tsconfig. Import from "@the-assembly/types" in both '
                  'server and client code. Use npm workspaces or turborepo if the project grows further.',
    },
    {
        'id': 'I-03',
        'title': 'Discord Activity: Verify Permissions-Policy Header Coverage',
        'severity': 'INFO',
        'category': 'Platform Compatibility',
        'file': 'server.ts (~L222)',
        'description': 'The Permissions-Policy middleware is applied after registerRoutes(), meaning it runs after '
                       'apiRoutes.ts middleware is registered. Express processes middleware in order, so API route '
                       'responses served entirely within registerRoutes() may not receive this header.',
        'action': 'Move the Permissions-Policy middleware to before the registerRoutes() call to ensure it applies universally.',
    },
]

add_issue_table(doc, issues)

# ─── SECTION 6: Refactor Sequence Recommendation ──────────────────────────────
add_heading(doc, '6. Recommended Refactor Sequence', 1, '1F3864')
doc.add_paragraph(
    'Given the interdependencies between the three monolithic files, the following sequence minimises regression risk:'
)
steps = [
    ('Phase 0 — Test Coverage First', 
     'Write unit tests for pure functions BEFORE touching any monolith. Target: computeEloChange, agenda evaluators, '
     'getExecutiveAction, deck utilities. These tests are the safety net for all subsequent refactoring.'),
    ('Phase 1 — Harden Critical Issues (C-01, C-02, H-01, H-03)',
     'Fix the four highest-risk issues in-place within the existing monoliths. These are surgical fixes that do not '
     'require architectural changes and should be done before any structural work begins.'),
    ('Phase 2 — Extract Service Modules from GameEngine',
     'Begin with statsService.ts (ELO, XP, match saving) — it has the clearest interface. '
     'Then timerService.ts, then aiService.ts, then redisService.ts. '
     'Keep GameEngine as a facade; do not change the public API surface until all services are extracted.'),
    ('Phase 3 — Split apiRoutes.ts into Domain Routers',
     'Create express.Router() modules for auth, profile, shop, admin, social. '
     'This is relatively low risk as each route is largely independent.'),
    ('Phase 4 — Slim Down server.ts',
     'Extract socket handlers, middleware/security, proxy, and webhook into dedicated files. '
     'server.ts should become a ~100-line orchestration file only.'),
    ('Phase 5 — Resend Integration',
     'Once DNS is verified, swap Nodemailer for Resend in the auth router. '
     'Add RESEND_API_KEY to Cloud Run secrets. Test with a staging email address before promoting to production.'),
]
for i, (title, body) in enumerate(steps, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}. {title}')
    r.bold = True
    r.font.size = Pt(10.5)
    doc.add_paragraph(body)

# ─── SECTION 7: Summary Metrics ───────────────────────────────────────────────
add_heading(doc, '7. Summary Metrics', 1, '1F3864')
table2 = doc.add_table(rows=8, cols=2)
table2.style = 'Table Grid'
rows_data = [
    ('Metric', 'Value'),
    ('Total lines of server-side code', '~7,500'),
    ('Total lines of client-side code', '~15,000 (est.)'),
    ('Critical Issues', '2'),
    ('High Priority Issues', '6'),
    ('Medium Priority Issues', '8 (incl. 3 planned refactors)'),
    ('Low / Informational Issues', '7'),
    ('Known Security Vulnerabilities (active)', '0'),
]
for i, (k, v) in enumerate(rows_data):
    cells = table2.rows[i].cells
    cells[0].paragraphs[0].clear()
    cells[1].paragraphs[0].clear()
    rk = cells[0].paragraphs[0].add_run(k)
    rv = cells[1].paragraphs[0].add_run(v)
    rk.font.size = Pt(9)
    rv.font.size = Pt(9)
    if i == 0:
        rk.bold = True
        rv.bold = True
        set_cell_bg(cells[0], '1F3864')
        set_cell_bg(cells[1], '1F3864')
        rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rv.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()
doc.add_paragraph(
    'Note: The three planned refactor items (M-05, M-06, M-07) are classified as MEDIUM rather than HIGH '
    'because the current monolithic code is functional and correct — the refactor is about long-term maintainability, '
    'not an urgent defect. It should, however, be completed before the next major feature addition.'
)

# Footer
doc.add_paragraph()
footer_p = doc.add_paragraph('The Assembly — Audit Report v1.0 | Horsemen Interactive | March 2026')
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer_p.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True

# Save
out_path = r'c:\Users\aeryt\Documents\The Assembly\Dev\main\the-assembly\The_Assembly_Audit_Report.docx'
doc.save(out_path)
print(f'Report saved to: {out_path}')
